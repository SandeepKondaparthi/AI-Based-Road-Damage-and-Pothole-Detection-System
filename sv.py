import os
from docx import Document

def is_code_file(file_name):
    """Check if the file has a code-related extension and is not minified or inside node_modules."""
    code_extensions = {".py", ".js", ".java", ".cpp", ".c", ".html", ".css", ".ts",".tsx",".jsx", ".kt", ".swift", ".php"}
    
    # Exclude minified files (e.g., .min.js, .min.css)
    if file_name.endswith(".min.js") or file_name.endswith(".min.css"):
        return False

    return any(file_name.endswith(ext) for ext in code_extensions)

def format_filename_as_heading(filename):
    """Formats a filename into a more readable heading."""
    name_without_ext = os.path.splitext(filename)[0]  # Remove extension
    return name_without_ext.replace("_", " ").replace("-", " ").title()

def count_code_files(source_folder):
    """Counts the total number of code files in the project directory, excluding node_modules and minified files."""
    count = 0
    for root, _, files in os.walk(source_folder):
        if "node_modules" in root:  # Skip node_modules
            continue
        count += sum(1 for file in files if is_code_file(file))
    return count

def copy_code_to_docx(source_folder, output_file):
    """Scans a project directory, extracts code, and saves it to a .docx file in a structured format.
       Stops execution if an error occurs in any file.
    """
    doc = Document()
    doc.add_heading('Project Code Compilation', level=1)
    
    # Get total code file count
    total_files = count_code_files(source_folder)
    doc.add_paragraph(f"📌 Total Code Files: {total_files}\n", style="Normal")

    file_count = 0  # Track processed files

    for root, _, files in os.walk(source_folder):
        if "node_modules" in root:  # Skip node_modules folder
            continue

        for file in sorted(files):  # Ensuring sequential order
            if is_code_file(file):
                file_path = os.path.join(root, file)
                try:
                    with open(file_path, "r", encoding="utf-8") as f:
                        code = f.read()

                    file_count += 1

                    # Add File Heading (Formatted)
                    doc.add_heading(f"{file_count}. {format_filename_as_heading(file)}", level=2)
                    doc.add_paragraph(f"📂 File Path: {file_path}\n", style="Normal")
                    
                    # Add Code Block
                    code_paragraph = doc.add_paragraph()
                    code_paragraph.add_run(code).bold = False
                    
                    # Add Separator for Better Readability
                    doc.add_paragraph("\n" + "=" * 80 + "\n")

                except Exception as e:
                    print(f"🚨 The code has stopped! Error occurred at file: {file_path}")
                    print(f"⚠️ Error: {e}")
                    return  # Stop execution immediately

    doc.save(output_file)
    print(f"✅ Code successfully copied to: {output_file}")
    print(f"📌 Total Code Files Processed: {file_count}")

if __name__ == "__main__":
    # Get user-defined project directory
    project_directory = input("Enter the full path of the project folder: ").strip()
    output_directory = input("Enter the full path where the output file should be saved: ").strip()

    # Validate the project folder
    if not os.path.exists(project_directory):
        print("❌ Error: The specified project folder does not exist.")
    elif not os.path.exists(output_directory):
        print("❌ Error: The specified output directory does not exist.")
    else:
        output_docx = os.path.join(output_directory, "Project_Code_Documentation.docx")
        copy_code_to_docx(project_directory, output_docx)
